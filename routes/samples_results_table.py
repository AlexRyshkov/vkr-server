from flask_restful import Resource
from flask import request, jsonify
import pandas as pd
import scipy.stats
import numpy as np
import json

from helpers import filter_by_normal_values, standardize_tests_results, get_filters_description
from queries import get_samples_results, get_test_name

from docx import Document


class SamplesResultsTable(Resource):
    def post(self):
        data = request.json
        filter = data['filter'] if 'filter' in data else None
        test_ids = list(map(lambda x: x, data['testIds']))
        samples_results = get_samples_results(test_ids=test_ids, filters=filter, have_all_tests=True)
        df = pd.DataFrame(samples_results)
        table_data = get_table_data(df, test_ids)
        table_data['filter_description'] = get_filters_description(filter)
        return jsonify(table_data)


class SamplesResultsTableDownload(Resource):
    def post(self):
        data = request.json
        filter = data['filter'] if 'filter' in data else None
        test_ids = list(map(lambda x: x, data['testIds']))
        samples_results = get_samples_results(test_ids=test_ids, filters=filter, have_all_tests=True)
        df = pd.DataFrame(samples_results)
        table_data = get_table_data(df, test_ids)
        test_names = list(map(lambda x: get_test_name(x), test_ids))
        save_table_data(table_data, "stats.docx", test_names, filter)
        return {
            "directory": "",
            "filename": "stats.docx"
        }


def get_table_data(df, test_ids):
    tableNormalValues = []
    df_normal_values = df.copy()
    df_normal_values = standardize_tests_results(df_normal_values)
    df_normal_values['Значение показателя'] = df_normal_values['result'].apply(
        lambda x: 'ниже нормы' if x < 0 else 'выше нормы' if x > 1 else 'норма')
    for test_id in test_ids:
        test_normal_values = df_normal_values.loc[df_normal_values['test_id'] == test_id]
        test_normal_values = test_normal_values.groupby(by=['Значение показателя']).result.count()
        d = test_normal_values.to_dict()
        d['Показатель'] = get_test_name(test_id)
        tableNormalValues.append(d)
    df['result'] = df['result'].astype(float)
    df = df.groupby(by=['test_id']).result.describe()
    df = df.round(3)
    df.insert(0, 'Показатель', list(map(lambda x: get_test_name(x), df.index.values)))
    for row in tableNormalValues:
        for column in ['Показатель', 'ниже нормы', 'норма', 'выше нормы']:
            if column not in row:
                row[column] = 0
    return {
        "tableStats": {
            "name": "Основные статистические показатели",
            "columns": list(df.columns),
            "values": json.loads(df.to_json(orient='records'))
        },
        "tableNormalValues": {
            "name": "Значения биохимических показателей",
            "columns": ['Показатель', 'ниже нормы', 'норма', 'выше нормы'],
            "values": tableNormalValues
        }
    }


def save_table_data(table_data, filename, test_names, filter):
    filter_description = get_filters_description(filter)
    document = Document()
    document.add_heading('Статистистика по биохимическим показателям', 0)
    document.add_paragraph(f'Выбранные биохимические показатели:{", ".join(test_names)}')
    document.add_paragraph(f'Фильтры:\r\n{filter_description}')
    for key in table_data:
        data = table_data[key]
        name = data['name']
        columns = data['columns']
        values = data['values']
        document.add_heading(name, 2)
        t = document.add_table(len(values) + 1, len(columns))
        t.style = 'Table Grid'
        for j in range(len(columns)):
            t.cell(0, j).text = columns[j]
        for i in range(len(values)):
            for j in range(len(columns)):
                column_name = columns[j]
                value = values[i][column_name]
                if isinstance(value, float):
                    t.cell(i + 1, j).text = str(round(format_number(value), 2))
                else:
                    t.cell(i + 1, j).text = str(value)
    document.save('stats.docx')


format_number = lambda n: n if n % 1 else int(n)
